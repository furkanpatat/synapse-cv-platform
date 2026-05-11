"""Extract raw text from PDF / DOCX files."""

import io
from pathlib import Path

import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    out: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            out.append(text)
    return "\n".join(out).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    out: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            out.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out).strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    if ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    raise ValueError(f"Unsupported file extension: {ext}")
